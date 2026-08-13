"""文档解析服务模块 - 从各类文件中提取纯文本

职责：文件 → 纯文本（文本提取层）。
与 document_splitter_service（纯文本 → 分片）职责分离：
- 本模块只负责"把各种格式的文件变成一段字符串"；
- 分片器只负责"把一段字符串切成适合向量化的块"。

当前支持的格式（通过文件扩展名自动分派到对应处理器）：
    .md   Markdown（直接读 UTF-8 文本）
    .txt  纯文本（直接读 UTF-8 文本）
    .pdf  PDF（pypdf 逐页提取）
    .docx Word 文档（python-docx 遍历段落）
老版 .doc 格式不被支持（python-docx 只认 .docx），上传时会由上层明确拒绝。
"""

from pathlib import Path
from typing import Dict

from loguru import logger

# 支持的扩展名集合（带点，小写）——"支持哪些格式"的唯一事实来源。
# 上传路由(file.py)的白名单校验、目录索引(vector_index_service)的遍历
# 都从这个常量 import，避免三处各自维护一份名单导致漏改。
SUPPORTED_EXTENSIONS = frozenset({".md", ".txt", ".pdf", ".docx"})


class DocumentParserService:
    """文档解析服务 - 按文件扩展名分派到对应的文本提取处理器"""

    def extract_text(self, file_path: str) -> str:
        """从文件中提取纯文本

        根据文件扩展名自动匹配对应的处理器：
        .md/.txt → 直接读文本；.pdf → pypdf 逐页提取；.docx → python-docx 遍历段落。

        Args:
            file_path: 文件路径

        Returns:
            str: 提取出的纯文本（可能为空字符串，如空文档/图片型 PDF）

        Raises:
            ValueError: 扩展名不在 SUPPORTED_EXTENSIONS 中
            RuntimeError: 文件存在但解析失败（损坏、加密无法解密等）
        """
        path = Path(file_path)
        ext = path.suffix.lower()
        handler_name = self._dispatcher.get(ext)
        if handler_name is None:
            raise ValueError(f"不支持的文件类型: {ext}")

        resolved = str(path.resolve())
        logger.info(f"解析文档: {path.name} (类型: {ext})")
        # 从方法名解析出真正的处理函数（避免类体内方法定义顺序问题）
        handler = getattr(self, handler_name)
        return handler(resolved)

    # ------- 处理器分派表 -------
    # 新增格式时：实现一个 _extract_xxx 方法 + 在这里登记一行（值为方法名字符串）即可。
    _dispatcher: Dict[str, str] = {
        ".md": "_extract_plain",
        ".txt": "_extract_plain",
        ".pdf": "_extract_pdf",
        ".docx": "_extract_docx",
    }

    # ------- 各格式的文本提取实现 -------

    def _extract_plain(self, file_path: str) -> str:
        """.md / .txt：直接按 UTF-8 读取纯文本"""
        path = Path(file_path)
        try:
            content = path.read_text(encoding="utf-8")
            logger.info(f"读取文本文件: {path.name}, 长度: {len(content)} 字符")
            return content
        except Exception as e:
            logger.error(f"读取文本文件失败: {path}, 错误: {e}")
            raise RuntimeError(f"解析失败 {path}: {e}") from e

    def _extract_pdf(self, file_path: str) -> str:
        """.pdf：使用 pypdf 逐页提取文本

        原理：PDF 内部每页是一个"内容流"（可含文本、图片、矢量），
        pypdf 解析内容流，把可提取的文字（文本层）按页取出来。
        注意：扫描件/图片型 PDF 没有文本层，extract_text() 会返回空串，
        需要 OCR 才能处理，属于已知限制。
        """
        from pypdf import PdfReader

        path = Path(file_path)
        try:
            reader = PdfReader(file_path)

            # 加密 PDF：先尝试用空密码解密（很多"已加密"文件其实无密码限制）
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception:
                    pass
                # decrypt 后仍未解除加密，说明有真实密码保护
                if reader.is_encrypted:
                    raise RuntimeError(
                        f"PDF 已加密，无法解析: {path}。请提供未加密版本（如去除密码后重新导出）"
                    )

            pages = []
            for page in reader.pages:
                text = page.extract_text()
                # 过滤无文本的页（空页/纯图片页），避免产生多余空行
                if text and text.strip():
                    pages.append(text.strip())

            content = "\n".join(pages)

            if not content:
                # 无文本层：很可能是扫描件/图片型 PDF
                logger.warning(
                    f"PDF 未提取到文本，可能为扫描件/图片型: {path}"
                )
                return ""

            logger.info(f"PDF 解析成功: {path.name}, 页数: {len(reader.pages)}, "
                        f"提取文本长度: {len(content)} 字符")
            return content

        except RuntimeError:
            raise
        except Exception as e:
            logger.error(f"PDF 解析失败: {path}, 错误: {e}")
            raise RuntimeError(f"解析失败 {path}: {e}") from e

    def _extract_docx(self, file_path: str) -> str:
        """.docx：使用 python-docx 遍历段落提取文本

        原理：Word 文档本质是 zip 包，内部是 XML 结构（document.xml）。
        python-docx 解析该 XML，通过 doc.paragraphs 暴露文档中的段落，
        每个段落 p.text 是这段文字的纯文本（去掉了样式信息）。
        表格本期不提取（保持简单），后续如需可扩展遍历 doc.tables。
        """
        from docx import Document

        path = Path(file_path)
        try:
            doc = Document(file_path)

            # 只提取非空段落，段落间用换行分隔（还原阅读顺序）
            paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            content = "\n".join(paragraphs)

            if not content:
                logger.warning(f"Word 文档未提取到段落文本: {path}")
                return ""

            logger.info(f"Word 解析成功: {path.name}, 段落数: {len(paragraphs)}, "
                        f"提取文本长度: {len(content)} 字符")
            return content

        except Exception as e:
            logger.error(f"Word 解析失败: {path}, 错误: {e}")
            raise RuntimeError(f"解析失败 {path}: {e}") from e


# 全局单例（与项目其他 service 保持一致）
document_parser_service = DocumentParserService()
